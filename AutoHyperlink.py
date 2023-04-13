import docx
import glob
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    r = paragraph.add_run ()
    r._r.append (hyperlink)

    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


files = glob.glob('*.txt')
for i in range(len(files)):
    s = str(i) + ": " + files[i]
    print(s)
n = int(input("Выберите файл ввода: "))
print("0: Консоль")
print("1: Текстовый файл")
print("2: Word")
m = int(input("Выберите формат вывода: "))

doc = docx.Document()
par = doc.add_paragraph()

file_input = open(files[n], "r", encoding="utf8")

if m == 1:
    file_output = open("Ссылки.txt", "w")

while True:
    line = file_input.readline()

    if not line:
        break

    s = line.split(' ')
    url = "https://yandex.ru/search/?text="

    for i in range(len(s)):
        url+=s[i]+"+"

    url = url[:-1] + '\n'

    if m == 0:
        print(url)
    if m == 1:
        file_output.write(url)
    if m == 2:
        add_hyperlink(par, line, url)

if (m == 2):
    doc.save('Ссылки.docx')
if (m == 0):
    input("Нажмите любую кнопку...")






